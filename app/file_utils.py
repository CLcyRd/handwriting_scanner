from pathlib import Path
from typing import List

import fitz  # PyMuPDF

# Ensure we have the correct fitz module from PyMuPDF
if not hasattr(fitz, "open"):
    try:
        import pymupdf
        fitz = pymupdf
    except ImportError:
        pass

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.models import TaskState


def convert_pdf_to_images(pdf_path: Path, output_dir: Path) -> List[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    image_paths: List[Path] = []
    with fitz.open(pdf_path) as doc:
        for index, page in enumerate(doc, start=1):
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image_path = output_dir / f"page_{index:04d}.png"
            pixmap.save(image_path)
            image_paths.append(image_path)
    return image_paths


def add_page_number(paragraph):
    """
    Append a PAGE field to the paragraph.
    """
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def build_docx(task: TaskState) -> Path:
    document = Document()
    ordered_pages = sorted(task.results.keys())
    
    for idx, page_number in enumerate(ordered_pages):
        # Create a new section for each page (except the first one which already exists)
        if idx > 0:
            document.add_section(WD_SECTION.NEW_PAGE)
        
        # Add content to the current section
        body = document.add_paragraph(task.results[page_number])
        body.paragraph_format.space_after = 0
        
        # Configure the footer for the current section
        section = document.sections[-1]
        section.footer.is_linked_to_previous = False
        
        # Clear any existing paragraphs in the footer (default footer might have one empty paragraph)
        for p in section.footer.paragraphs:
            p._element.getparent().remove(p._element)
            
        footer_para = section.footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add "Word Page: " + PAGE field + " | PDF Page: " + page_number
        footer_para.add_run("Word页码: ")
        add_page_number(footer_para)
        footer_para.add_run(f" | 内容在pdf中页码: {page_number}")

    output_path = task.pdf_path.parent / f"{task.pdf_path.stem}_recognized.docx"
    document.save(output_path)
    return output_path
